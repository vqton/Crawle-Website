
from bs4 import BeautifulSoup
import requests
import docx

hdr = {
    'User-Agent': 'Mozilla/5.0'}
url = 'https://tuoinung.com/2020/10/truyen-sex-chi-toi-va-co-thu-ky-cua-bo-toi.html'
html = requests.get(url, headers=hdr).content
soup = BeautifulSoup(html, 'html.parser')
# print(i.text)
part = 1
mydoc = docx.Document()
urls = []
urls.append(url)

sExclusive = "Truyện sex hay với đầy đủ các thể loại:"
next_page = soup.find("a", text="Đọc Tiếp").get('href')
while next_page:
    url = next_page
    html = requests.get(url, headers=hdr).content
    soup = BeautifulSoup(html, 'html.parser')
    try:
        next_page = soup.find("a", text="Đọc Tiếp").get('href')
        urls.append(url)

    except:
        break

part = 1
for url1 in urls:

    html = requests.get(url1, headers=hdr).content
    soup = BeautifulSoup(html, 'html.parser')
    title = soup.find("h1").text
    content = soup.select("div > p > span")

    last_item = len(content) - 1
    for i in content:

        if content.index(i) == 0:
            mydoc.add_heading('Phần ' + str(part), level=1)
            mydoc.save('a.docx')

        if last_item == content.index(i):
            mydoc.add_paragraph("Kết thúc " + i.text)
            s = "="
            for r in range(1, 80):
                s += "="
            mydoc.add_paragraph(s)
            mydoc.save('a.docx')
        else:
            mydoc.add_paragraph(i.text)
            mydoc.save('a.docx')

    part = part + 1

    # content = soup.find("div", class_="entry-content")
    # for i in content.find_all("p"):
    #     print(i.text)
# def getContent(url):
#     soup = BeautifulSoup(url, 'html.parser')
#     file_name = soup.h1.string.replace(" Truyện Sex: ", "")

#     content = soup.select("div > p > span")
#     for i in content:
#         f = open(f'{file_name}_P{part}.txt', "a", encoding="utf-8")
#         f.write(i.text + " \n")
#         f.close()


# # print(soup.prettify())
# next_page = soup.find("a", text="Đọc Tiếp")
# while next_page is not None:
#     url1 = next_page.get('href')
#     getContent(url1)
# # print(soup.h1.string)
# # content = soup.select("div > p > span")

# print(type(content))

# for i in content.find_all('p'):
#     a = i.find('span')
#     print(a.text)

# for i in content:
#     print(i.p.span)
# print(i.text.rstrip())

# new_feed = soup.find('div', class_='box').find('p')
# print(type(new_feed))
